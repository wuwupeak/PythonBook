#2c5e80d1-0c14-463c-9440-2dc889d195d7

import os

#必须安装的docx兼容包
# pip install ./pythonstudy/python_docx-0.8.7-py2.py3-none-any.whl
import docx  # pip install python-docx
from docx.oxml.ns import qn
import re

codeFileABS = os.path.abspath(__file__) # 获取处理代码文件所在的绝对路径信息
print(codeFileABS)
codeFileABSDir = codeFileABS.replace('\\WriteToWord.py','') # 获取处理代码文件夹所在的绝对路径信息
print(codeFileABSDir)
docxSaveDir = codeFileABS.replace('WriteToWord.py','') # 生成处理结果docx存储需要的绝对路径
folders = os.listdir(codeFileABSDir) # 获取获取处理代码文件夹所在的绝对路径下的文件夹列表
print(folders)

doc = docx.Document()
doc.styles['Normal'].font.name = u'微软雅黑'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
oneClassCount = 1

# paragraphIndex = -1 #
for folder in folders:
        if (folder == "000入门部分") or (folder == '010基础部分') or (folder == '020进阶部分') or (folder == '030高级部分'):  #只对要进行写入的文件夹进行筛选
                folderPathABS = os.path.join(codeFileABSDir, folder)
                print (folderPathABS)
                subFolders = os.listdir(folderPathABS)

                oneClassPrefix = str(oneClassCount)
                doc.add_heading(oneClassPrefix + ' ' + folder[3:], level=1) #使用一级标题样式
                oneClassCount = oneClassCount + 1
                twoClassCount = 1
                for subFolder in subFolders:                  
                        
                        doc.add_page_break() # 插入换页 
                        
                        subFolderPathABS = os.path.join(folderPathABS,subFolder)
                        print(subFolderPathABS)

                        twoClassPrefix = oneClassPrefix + '.' +str(twoClassCount)
                        doc.add_heading(twoClassPrefix + ' ' + subFolder[3:], level=2) #使用二级标题样式
                        twoClassCount = twoClassCount + 1

                        files = os.listdir(subFolderPathABS)
                        threeClassCount = 1
                        for fileOne in files:                     
                                filePathABS = os.path.join(subFolderPathABS,fileOne)
                                print(filePathABS)
                                if (os.path.isfile(filePathABS)):
                                        fileNameSplit = os.path.splitext(filePathABS)
                                        if(fileNameSplit[1] == '.py'):
                                                #(^[A-Za-z]*)匹配0个或者多个字母
                                                #(\()匹配右括号
                                                #(\d{2,3})匹配两到三位的数字
                                                wordRE = re.compile(r"(^[A-Za-z]*)(\()(\d{2,3})")
                                                matchObject = re.match(wordRE,fileOne)

                                                if matchObject:
                                                        number = matchObject.group(3) #获得数字匹配的部分字符串
                                                        fileOne = fileOne.replace(number,'') #去除数字部分字符串

                                                threeClassPrefix = twoClassPrefix + '.' + str(threeClassCount)
                                                doc.add_heading(threeClassPrefix + ' ' + fileOne[0:-3], level=3) #使用三级标题样式
                                                threeClassCount = threeClassCount + 1

                                                with open(filePathABS,encoding='UTF-8') as fileObject: #转义符'\\'和读取中文的编码'UTF-8'
                                                        contents = fileObject.read()
                                                        doc.add_paragraph(contents.rstrip())                                                       
                                                        print(contents.rstrip())                                               
doc.save(os.path.join(docxSaveDir,'coding.docx'))
